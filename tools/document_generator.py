"""
文档生成工具
支持生成会议纪要、工作总结、分析报告等 Word 文档
集成 LLM 智能内容生成能力
"""
import os
import json
import re
from datetime import datetime
from typing import Optional, List, Type, Dict, Any
from pydantic import BaseModel, Field
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging

from config.settings import settings
from .base import BaseTool

logger = logging.getLogger(__name__)


class DocumentGeneratorInput(BaseModel):
    """文档生成工具输入参数模式"""
    doc_type: str = Field(description="文档类型: meeting_notes, work_summary, report")
    title: Optional[str] = Field(None, description="文档标题")
    content: Optional[str] = Field(None, description="文档主要内容（可由LLM自动生成）")
    auto_generate: bool = Field(False, description="是否启用LLM自动生成内容")
    context: Optional[str] = Field(None, description="上下文信息，用于LLM生成内容")
    participants: Optional[List[str]] = Field(None, description="参会人员列表（会议纪要）")
    agenda: Optional[List[str]] = Field(None, description="会议议程（会议纪要）")
    achievements: Optional[List[str]] = Field(None, description="工作成果（工作总结）")
    challenges: Optional[List[str]] = Field(None, description="遇到的挑战（工作总结）")
    plans: Optional[List[str]] = Field(None, description="未来计划（工作总结）")
    data_summary: Optional[str] = Field(None, description="数据分析摘要（报告）")


class DocumentGeneratorTool(BaseTool):
    """
    文档生成工具

    功能：
    1. 生成会议纪要（包含参会人员、议程、决议等）
    2. 生成工作总结（周报/月报/年报）
    3. 生成分析报告（结合数据分析结果）
    4. 支持 LLM 智能内容生成
    """

    name: str = "document_generator"
    description: str = """
    文档生成工具，用于创建专业的 Word 文档。
    支持的文档类型：
    - meeting_notes: 会议纪要，需要提供 title、participants、agenda、content
    - work_summary: 工作总结，需要提供 title、achievements、challenges、plans
    - report: 分析报告，需要提供 title、data_summary
    支持 auto_generate=true 启用 LLM 自动生成内容。
    生成的文档保存在 data/documents 目录下。
    """
    args_schema: Type[BaseModel] = DocumentGeneratorInput

    _llm = None

    def __init__(self, llm=None, **kwargs):
        super().__init__(**kwargs)
        if llm:
            self._llm = llm
        elif settings.DASHSCOPE_API_KEY:
            import os as os_module
            os_module.environ["DASHSCOPE_API_KEY"] = settings.DASHSCOPE_API_KEY
            from langchain_community.llms import Tongyi
            self._llm = Tongyi(model_name=settings.LLM_MODEL_NAME)

    def _generate_content_with_llm(self, doc_type: str, context: str) -> Dict[str, Any]:
        """使用 LLM 智能生成文档内容"""
        if not self._llm:
            logger.warning("LLM 未初始化，无法自动生成内容")
            return {}

        prompts = {
            "meeting_notes": f"""根据以下会议信息，生成会议纪要的内容要点。请以JSON格式返回，包含：
- content: 会议主要内容摘要（200字以内）
- decisions: 决议事项列表（3-5条）
- action_items: 待办事项列表（3-5条）

会议信息：
{context}

请只返回JSON，不要包含其他文字。""",

            "work_summary": f"""根据以下工作信息，生成工作总结的内容要点。请以JSON格式返回，包含：
- achievements: 工作成果列表（3-5条）
- challenges: 遇到的挑战（2-3条）
- solutions: 解决方案（2-3条）
- plans: 下阶段计划（3-5条）

工作信息：
{context}

请只返回JSON，不要包含其他文字。""",

            "report": f"""根据以下分析数据，生成分析报告的结论与建议。请以JSON格式返回，包含：
- summary: 报告摘要（100字以内）
- conclusions: 主要结论（3-5条）
- recommendations: 建议措施（3-5条）

分析数据：
{context}

请只返回JSON，不要包含其他文字。"""
        }

        prompt = prompts.get(doc_type, prompts["report"])

        try:
            response = self._llm.invoke(prompt)
            match = re.search(r'\{[^{}]*\}', response, re.DOTALL)
            if match:
                return json.loads(match.group())
        except Exception as e:
            logger.error(f"LLM 内容生成失败: {e}")

        return {}

    def _execute(self, doc_type: str, title: Optional[str] = None,
                 content: Optional[str] = None, auto_generate: bool = False,
                 context: Optional[str] = None, **kwargs) -> str:
        """生成文档"""
        if doc_type == "meeting_notes":
            participants = kwargs.get("participants")
            agenda = kwargs.get("agenda")
            return self._generate_meeting_notes(title, content, auto_generate, context, participants, agenda)
        elif doc_type == "work_summary":
            achievements = kwargs.get("achievements")
            challenges = kwargs.get("challenges")
            plans = kwargs.get("plans")
            return self._generate_work_summary(title, auto_generate, context, achievements, challenges, plans)
        elif doc_type == "report":
            data_summary = kwargs.get("data_summary")
            return self._generate_report(title, content, auto_generate, context, data_summary)
        else:
            return f"未知文档类型: {doc_type}，支持: meeting_notes, work_summary, report"

    def _generate_meeting_notes(self, title: Optional[str], content: Optional[str],
                                auto_generate: bool, context: Optional[str],
                                participants: Optional[List[str]] = None,
                                agenda: Optional[List[str]] = None) -> str:
        """生成会议纪要"""
        try:
            generated_content = {}
            if auto_generate and context:
                generated_content = self._generate_content_with_llm("meeting_notes", context)

            doc = Document()

            title_text = title or "会议纪要"
            heading = doc.add_heading(title_text, level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_heading('会议信息', level=1)
            doc.add_paragraph(f'会议时间：{datetime.now().strftime("%Y年%m月%d日 %H:%M")}')

            if participants:
                doc.add_paragraph(f'参会人员：{", ".join(participants)}')
            elif generated_content.get("participants"):
                doc.add_paragraph(f'参会人员：{", ".join(generated_content["participants"])}')
            else:
                doc.add_paragraph('参会人员：待补充')

            if agenda:
                doc.add_heading('会议议程', level=1)
                for item in agenda:
                    doc.add_paragraph(item, style='List Bullet')

            doc.add_heading('会议内容', level=1)
            if content:
                for para in content.split('\n'):
                    if para.strip():
                        doc.add_paragraph(para.strip())
            elif generated_content.get("content"):
                doc.add_paragraph(generated_content["content"])
            else:
                doc.add_paragraph('会议内容待补充')

            doc.add_heading('决议与待办事项', level=1)
            decisions = generated_content.get("decisions", [])
            action_items = generated_content.get("action_items", [])

            if decisions:
                for item in decisions:
                    doc.add_paragraph(f'• {item}', style='List Bullet')
            else:
                doc.add_paragraph('• 待补充', style='List Bullet')

            if action_items:
                doc.add_paragraph('待办事项：', style='List Bullet')
                for item in action_items:
                    doc.add_paragraph(f'  - {item}', style='List Bullet')

            file_path = self._save_document(doc, f"meeting_notes_{datetime.now().strftime('%Y%m%d_%H%M%S')}")
            return f"会议纪要生成成功！文件保存至: {file_path}"
        except Exception as e:
            logger.error(f"生成会议纪要失败: {e}")
            return f"生成会议纪要失败: {str(e)}"

    def _generate_work_summary(self, title: Optional[str], auto_generate: bool,
                               context: Optional[str],
                               achievements: Optional[List[str]] = None,
                               challenges: Optional[List[str]] = None,
                               plans: Optional[List[str]] = None) -> str:
        """生成工作总结"""
        try:
            generated_content = {}
            if auto_generate and context:
                generated_content = self._generate_content_with_llm("work_summary", context)

            doc = Document()

            title_text = title or f"{datetime.now().strftime('%Y年%m月')}工作总结"
            heading = doc.add_heading(title_text, level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph(f'汇报人：________  日期：{datetime.now().strftime("%Y年%m月%d日")}')
            doc.add_paragraph()

            doc.add_heading('一、工作成果', level=1)
            items = achievements or generated_content.get("achievements", [])
            if items:
                for item in items:
                    doc.add_paragraph(f'• {item}', style='List Bullet')
            else:
                doc.add_paragraph('• 本期工作成果待补充', style='List Bullet')

            doc.add_heading('二、遇到的挑战与解决方案', level=1)
            challenge_items = challenges or generated_content.get("challenges", [])
            solutions = generated_content.get("solutions", [])

            if challenge_items:
                for item in challenge_items:
                    doc.add_paragraph(f'• 挑战：{item}', style='List Bullet')
            if solutions:
                for item in solutions:
                    doc.add_paragraph(f'• 解决方案：{item}', style='List Bullet')
            if not challenge_items and not solutions:
                doc.add_paragraph('• 挑战与解决方案待补充', style='List Bullet')

            doc.add_heading('三、下阶段工作计划', level=1)
            plan_items = plans or generated_content.get("plans", [])
            if plan_items:
                for item in plan_items:
                    doc.add_paragraph(f'• {item}', style='List Bullet')
            else:
                doc.add_paragraph('• 下阶段计划待补充', style='List Bullet')

            file_path = self._save_document(doc, f"work_summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}")
            return f"工作总结生成成功！文件保存至: {file_path}"
        except Exception as e:
            logger.error(f"生成工作总结失败: {e}")
            return f"生成工作总结失败: {str(e)}"

    def _generate_report(self, title: Optional[str], content: Optional[str],
                        auto_generate: bool, context: Optional[str],
                        data_summary: Optional[str] = None) -> str:
        """生成分析报告"""
        try:
            generated_content = {}
            if auto_generate and context:
                generated_content = self._generate_content_with_llm("report", context)

            doc = Document()

            title_text = title or "数据分析报告"
            heading = doc.add_heading(title_text, level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph(f'生成时间：{datetime.now().strftime("%Y年%m月%d日 %H:%M")}')
            doc.add_paragraph()

            doc.add_heading('报告摘要', level=1)
            summary = data_summary or generated_content.get("summary")
            if summary:
                doc.add_paragraph(summary)
            else:
                doc.add_paragraph('本报告基于数据分析工具生成，详细内容如下。')

            doc.add_heading('详细分析', level=1)
            if content:
                for para in content.split('\n'):
                    if para.strip():
                        doc.add_paragraph(para.strip())
            else:
                doc.add_paragraph('详细分析内容待补充。')

            doc.add_heading('结论与建议', level=1)
            doc.add_paragraph('基于以上分析，提出以下建议：')

            conclusions = generated_content.get("conclusions", [])
            recommendations = generated_content.get("recommendations", [])

            if conclusions:
                doc.add_paragraph('主要结论：', style='List Bullet')
                for item in conclusions:
                    doc.add_paragraph(f'  - {item}', style='List Bullet')

            if recommendations:
                doc.add_paragraph('建议措施：', style='List Bullet')
                for item in recommendations:
                    doc.add_paragraph(f'  - {item}', style='List Bullet')

            if not conclusions and not recommendations:
                doc.add_paragraph('• 建议待补充', style='List Bullet')

            file_path = self._save_document(doc, f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}")
            return f"分析报告生成成功！文件保存至: {file_path}"
        except Exception as e:
            logger.error(f"生成报告失败: {e}")
            return f"生成报告失败: {str(e)}"

    def _save_document(self, doc: Document, base_name: str) -> str:
        """保存文档到指定目录"""
        output_dir = settings.DOCUMENT_OUTPUT_DIR
        os.makedirs(output_dir, exist_ok=True)

        file_name = f"{base_name}.docx"
        file_path = os.path.join(output_dir, file_name)
        doc.save(file_path)
        return file_path
