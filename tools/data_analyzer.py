"""
数据分析工具
支持读取 Excel/CSV、数据清洗、统计分析、可视化
自动生成分析报告并集成文档生成功能
"""
import os
import json
import pandas as pd
import matplotlib.pyplot as plt
from typing import Optional, Dict, Any, Type
from pydantic import BaseModel, Field
import logging

from config.settings import settings
from .base import BaseTool

logger = logging.getLogger(__name__)

plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'DejaVu Sans']
plt.rcParams['axes.unicode_minus'] = False


class DataAnalyzerInput(BaseModel):
    """数据分析工具输入参数模式"""
    action: str = Field(description="操作类型: read_data, clean_data, analyze, visualize, generate_report")
    file_path: Optional[str] = Field(None, description="数据文件路径（Excel/CSV）")
    sheet_name: Optional[str] = Field(None, description="Excel工作表名称")
    session_id: Optional[str] = Field(None, description="会话ID，用于隔离多用户数据")
    auto_generate_doc: Optional[bool] = Field(False, description="是否自动生成Word报告")


class DataAnalyzerTool(BaseTool):
    """
    数据分析工具

    功能：
    1. 读取 Excel/CSV 数据
    2. 数据清洗（去重、处理缺失值）
    3. 统计分析（描述性统计、相关性）
    4. 数据可视化（生成图表）
    5. 自动生成 Word 分析报告
    """

    name: str = "data_analyzer"
    description: str = """
    数据分析工具，用于处理 Excel/CSV 文件。
    支持的操作：
    - read_data: 读取数据文件，需要提供 file_path
    - clean_data: 清洗已加载的数据（去重、删除缺失值）
    - analyze: 生成数据统计分析报告
    - visualize: 生成可视化图表
    - generate_report: 生成完整的 Word 分析报告（自动整合分析和可视化结果）
    注意：需要先执行 read_data 才能进行后续操作。
    会话级别的数据隔离：不同 session_id 的数据互不影响。
    """
    args_schema: Type[BaseModel] = DataAnalyzerInput

    _data_registry: Dict[str, pd.DataFrame] = {}
    _current_file_registry: Dict[str, str] = {}

    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self._session_id = None

    @property
    def _data(self) -> Optional[pd.DataFrame]:
        return self._data_registry.get(self._session_id or "default")

    @_data.setter
    def _data(self, value: pd.DataFrame):
        key = self._session_id or "default"
        if value is None:
            self._data_registry.pop(key, None)
        else:
            self._data_registry[key] = value

    @property
    def _current_file(self) -> Optional[str]:
        return self._current_file_registry.get(self._session_id or "default")

    @_current_file.setter
    def _current_file(self, value: str):
        key = self._session_id or "default"
        if value is None:
            self._current_file_registry.pop(key, None)
        else:
            self._current_file_registry[key] = value

    def _execute(self, action: str, file_path: Optional[str] = None,
                 sheet_name: Optional[str] = None, session_id: Optional[str] = None,
                 auto_generate_doc: bool = False, **kwargs) -> str:
        """
        执行数据分析操作
        """
        self._session_id = session_id or "default"

        if action == "read_data":
            return self._read_data(file_path, sheet_name)
        elif action == "clean_data":
            return self._clean_data()
        elif action == "analyze":
            return self._analyze()
        elif action == "visualize":
            return self._visualize()
        elif action == "generate_report":
            return self._generate_full_report(auto_generate_doc)
        else:
            return f"未知操作: {action}，支持的操作: read_data, clean_data, analyze, visualize, generate_report"

    def _get_analysis_summary(self) -> str:
        """获取分析结果的摘要，用于填充报告"""
        if self._data is None:
            return "无数据"

        df = self._data
        summary_parts = []

        summary_parts.append(f"数据规模: {len(df)} 行, {len(df.columns)} 列")
        summary_parts.append(f"列名: {', '.join(df.columns.astype(str))}")

        numeric_cols = df.select_dtypes(include=['number']).columns
        if len(numeric_cols) > 0:
            summary_parts.append(f"\n数值列 ({len(numeric_cols)} 个): {', '.join(numeric_cols)}")
            for col in numeric_cols:
                summary_parts.append(f"  - {col}: 均值={df[col].mean():.2f}, 最大={df[col].max():.2f}, 最小={df[col].min():.2f}")

        text_cols = df.select_dtypes(include=['object']).columns
        if len(text_cols) > 0:
            summary_parts.append(f"\n文本列 ({len(text_cols)} 个): {', '.join(text_cols)}")

        return "\n".join(summary_parts)

    def _generate_full_report(self, auto_generate_doc: bool = False) -> str:
        """生成完整分析报告"""
        if self._data is None:
            return "请先使用 read_data 读取数据文件"

        try:
            results = []
            results.append("=" * 50)
            results.append("数据分柝报告")
            results.append("=" * 50)

            df = self._data
            results.append(f"\n【数据概览】")
            results.append(f"总行数: {len(df)}")
            results.append(f"总列数: {len(df.columns)}")

            numeric_cols = df.select_dtypes(include=['number']).columns
            if len(numeric_cols) > 0:
                results.append(f"\n【数值分析】")
                stats = df[numeric_cols].describe()
                results.append(stats.to_string())

                if len(numeric_cols) > 1:
                    results.append(f"\n【相关性分析】")
                    corr = df[numeric_cols].corr().round(2)
                    results.append(corr.to_string())

            text_cols = df.select_dtypes(include=['object']).columns
            if len(text_cols) > 0:
                results.append(f"\n【分类统计】")
                for col in text_cols[:5]:
                    value_counts = df[col].value_counts().head(5)
                    results.append(f"{col}: {dict(value_counts)}")

            report_text = "\n".join(results)

            if auto_generate_doc:
                try:
                    from docx import Document
                    from docx.shared import Pt, RGBColor
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    from datetime import datetime

                    doc = Document()
                    doc.add_heading('数据分析报告', level=0)

                    doc.add_paragraph(f'生成时间: {datetime.now().strftime("%Y年%m月%d日 %H:%M")}')
                    doc.add_paragraph(f'数据文件: {self._current_file or "未知"}')
                    doc.add_paragraph(f'数据规模: {len(df)} 行 x {len(df.columns)} 列')

                    doc.add_heading('一、数据概览', level=1)
                    doc.add_paragraph(f'共 {len(df)} 条记录，{len(df.columns)} 个字段。')
                    doc.add_paragraph(f'字段列表: {", ".join(df.columns.astype(str))}')

                    if len(numeric_cols) > 0:
                        doc.add_heading('二、数值分析', level=1)
                        doc.add_paragraph('描述性统计:')
                        doc.add_paragraph(stats.to_string())

                        if len(numeric_cols) > 1:
                            doc.add_heading('三、相关性分析', level=1)
                            doc.add_paragraph('相关系数矩阵:')
                            doc.add_paragraph(corr.to_string())

                    doc.add_heading('四、分类统计', level=1)
                    for col in text_cols[:5]:
                        doc.add_paragraph(f'{col}分布:')
                        vc = df[col].value_counts().head(5)
                        for val, count in vc.items():
                            doc.add_paragraph(f'  {val}: {count}', style='List Bullet')

                    viz_dir = settings.VISUALIZATION_OUTPUT_DIR
                    heatmap_path = os.path.join(viz_dir, "correlation_heatmap.png")
                    if os.path.exists(heatmap_path):
                        doc.add_heading('五、可视化图表', level=1)
                        doc.add_picture(heatmap_path, width=6*1024*1024/1024)

                    output_dir = settings.DOCUMENT_OUTPUT_DIR
                    os.makedirs(output_dir, exist_ok=True)
                    file_name = f"数据分析报告_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                    file_path = os.path.join(output_dir, file_name)
                    doc.save(file_path)
                    report_text += f"\n\n✅ Word报告已生成: {file_path}"
                except ImportError:
                    report_text += "\n\n⚠️ 报告生成依赖 python-docx，请确保已安装"
                except Exception as e:
                    report_text += f"\n\n⚠️ Word报告生成失败: {str(e)}"

            return report_text

        except Exception as e:
            logger.error(f"生成报告失败: {e}")
            return f"生成报告失败: {str(e)}"
    
    def _read_data(self, file_path: Optional[str], sheet_name: Optional[str]) -> str:
        """读取数据文件"""
        if not file_path:
            file_path = settings.DEFAULT_DATA_PATH
        
        # 如果指定路径不存在，尝试在上传目录中查找
        original_path = file_path
        if not os.path.exists(file_path):
            upload_dir = "data/uploads"
            file_name = os.path.basename(file_path)
            upload_path = os.path.join(upload_dir, file_name)
            if os.path.exists(upload_path):
                file_path = upload_path
            else:
                return f"文件不存在: {original_path}"
        
        try:
            ext = os.path.splitext(file_path)[1].lower()
            if ext in ['.xlsx', '.xls']:
                data = pd.read_excel(file_path, sheet_name=sheet_name)
                # 如果返回的是字典（多个工作表），选择第一个工作表
                if isinstance(data, dict) and data:
                    self._data = list(data.values())[0]
                else:
                    self._data = data
            elif ext == '.csv':
                self._data = pd.read_csv(file_path)
            else:
                return f"不支持的文件格式: {ext}，请使用 Excel 或 CSV 文件"
            
            # 确保_data是DataFrame对象
            if not isinstance(self._data, pd.DataFrame):
                return f"读取数据失败: 数据格式不正确，预期DataFrame对象"
            
            self._current_file = file_path
            rows, cols = self._data.shape
            return f"成功读取数据文件: {file_path}\n数据规模: {rows} 行, {cols} 列\n列名: {', '.join(self._data.columns.astype(str))}"
        except Exception as e:
            logger.error(f"读取数据失败: {e}")
            return f"读取数据失败: {str(e)}"
    
    def _clean_data(self) -> str:
        """清洗数据"""
        if self._data is None:
            return "请先使用 read_data 读取数据文件"
        
        try:
            original_rows = len(self._data)
            
            # 删除重复行
            self._data = self._data.drop_duplicates()
            
            # 删除全为空的行
            self._data = self._data.dropna(how='all')
            
            # 对于数值列，用中位数填充缺失值
            numeric_cols = self._data.select_dtypes(include=['number']).columns
            for col in numeric_cols:
                # 跳过空列或全为 NaN 的列
                if self._data[col].notna().any():
                    median_val = self._data[col].median()
                    self._data[col] = self._data[col].fillna(median_val)
            
            # 对于文本列，用"未知"填充
            text_cols = self._data.select_dtypes(include=['object']).columns
            for col in text_cols:
                self._data[col] = self._data[col].fillna("未知")
            
            new_rows = len(self._data)
            removed = original_rows - new_rows
            return f"数据清洗完成。原始行数: {original_rows}，清洗后行数: {new_rows}，移除 {removed} 行无效数据。"
        except Exception as e:
            logger.error(f"数据清洗失败: {e}")
            return f"数据清洗失败: {str(e)}"
    
    def _analyze(self) -> str:
        """统计分析"""
        if self._data is None:
            return "请先使用 read_data 读取数据文件"
        
        try:
            df = self._data
            result_parts = []
            
            # 基本统计信息
            result_parts.append("=== 数据概览 ===")
            result_parts.append(f"总行数: {len(df)}")
            result_parts.append(f"总列数: {len(df.columns)}")
            result_parts.append(f"列名: {', '.join(df.columns.astype(str))}")
            
            # 数值列统计
            numeric_cols = df.select_dtypes(include=['number']).columns
            if len(numeric_cols) > 0:
                result_parts.append("\n=== 数值列统计 ===")
                stats = df[numeric_cols].describe().to_string()
                result_parts.append(stats)
                
                # 相关性分析（如果有多列）
                if len(numeric_cols) > 1:
                    result_parts.append("\n=== 相关性矩阵 ===")
                    corr = df[numeric_cols].corr().round(2).to_string()
                    result_parts.append(corr)
            
            # 文本列统计
            text_cols = df.select_dtypes(include=['object']).columns
            if len(text_cols) > 0:
                result_parts.append("\n=== 文本列统计 ===")
                for col in text_cols:
                    unique_count = df[col].nunique()
                    result_parts.append(f"{col}: 唯一值数量 {unique_count}")
            
            return "\n".join(result_parts)
        except Exception as e:
            logger.error(f"数据分析失败: {e}")
            return f"数据分析失败: {str(e)}"
    
    def _visualize(self) -> str:
        """生成可视化图表"""
        if self._data is None:
            return "请先使用 read_data 读取数据文件"
        
        try:
            df = self._data
            numeric_cols = df.select_dtypes(include=['number']).columns
            
            if len(numeric_cols) == 0:
                return "数据中没有数值列，无法生成可视化图表"
            
            output_dir = settings.VISUALIZATION_OUTPUT_DIR
            os.makedirs(output_dir, exist_ok=True)
            
            generated_files = []
            
            # 为每个数值列生成直方图
            for col in numeric_cols[:3]:  # 限制最多3个图表
                plt.figure(figsize=(10, 6))
                plt.hist(df[col].dropna(), bins=20, edgecolor='black', alpha=0.7)
                plt.title(f'{col} 分布直方图')
                plt.xlabel(col)
                plt.ylabel('频数')
                plt.grid(True, alpha=0.3)
                
                file_name = f"{col}_histogram.png"
                file_path = os.path.join(output_dir, file_name)
                plt.savefig(file_path, dpi=150, bbox_inches='tight')
                plt.close()
                generated_files.append(file_path)
            
            # 如果有多个数值列，生成相关性热力图
            if len(numeric_cols) > 1:
                plt.figure(figsize=(10, 8))
                corr = df[numeric_cols].corr()
                plt.imshow(corr, cmap='coolwarm', aspect='auto')
                plt.colorbar(label='相关系数')
                plt.xticks(range(len(corr.columns)), corr.columns, rotation=45, ha='right')
                plt.yticks(range(len(corr.columns)), corr.columns)
                plt.title('特征相关性热力图')
                
                file_path = os.path.join(output_dir, "correlation_heatmap.png")
                plt.savefig(file_path, dpi=150, bbox_inches='tight')
                plt.close()
                generated_files.append(file_path)
            
            return f"可视化图表生成成功，共生成 {len(generated_files)} 个图表:\n" + "\n".join(generated_files)
        except Exception as e:
            logger.error(f"可视化生成失败: {e}")
            return f"可视化生成失败: {str(e)}"