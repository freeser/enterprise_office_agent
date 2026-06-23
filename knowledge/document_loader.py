"""
通用文档加载器
支持 PDF、Word、Excel、TXT 等多种格式，自动进行文本分块
"""
import os
from typing import List, Optional
import logging

from langchain_core.documents import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
import fitz  # PyMuPDF
from docx import Document as DocxDocument
import pandas as pd

logger = logging.getLogger(__name__)


class UniversalDocumentLoader:
    """通用文档加载器"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """
        初始化加载器
        
        Args:
            chunk_size: 文本分块大小
            chunk_overlap: 分块重叠大小
        """
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", "。", "！", "？", ".", "!", "?", " ", ""],
            length_function=len
        )
    
    def load_file(self, file_path: str, metadata: Optional[dict] = None) -> List[Document]:
        """
        根据文件扩展名自动选择加载方式
        
        Args:
            file_path: 文件路径
            metadata: 额外的元数据
        
        Returns:
            Document 列表
        """
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.pdf':
            return self._load_pdf(file_path, metadata)
        elif ext in ['.docx', '.doc']:
            return self._load_docx(file_path, metadata)
        elif ext in ['.xlsx', '.xls']:
            return self._load_excel(file_path, metadata)
        elif ext == '.txt':
            return self._load_txt(file_path, metadata)
        else:
            logger.warning(f"不支持的文件格式: {ext}")
            return []
    
    def _load_pdf(self, file_path: str, metadata: Optional[dict] = None) -> List[Document]:
        """加载 PDF 文件"""
        documents = []
        base_metadata = metadata or {}
        base_metadata.update({"source": file_path, "file_type": "pdf"})
        
        try:
            doc = fitz.open(file_path)
            full_text = ""
            
            for page_num, page in enumerate(doc):
                page_text = page.get_text()
                if page_text.strip():
                    full_text += f"\n--- 第 {page_num + 1} 页 ---\n{page_text}"
            
            doc.close()
            
            # 分块
            chunks = self.text_splitter.split_text(full_text)
            for i, chunk in enumerate(chunks):
                meta = base_metadata.copy()
                meta["chunk_index"] = i
                documents.append(Document(page_content=chunk, metadata=meta))
            
            logger.info(f"PDF 加载完成: {file_path}, 共 {len(documents)} 个片段")
        except Exception as e:
            logger.error(f"PDF 加载失败 {file_path}: {e}")
        
        return documents
    
    def _load_docx(self, file_path: str, metadata: Optional[dict] = None) -> List[Document]:
        """加载 Word 文档"""
        documents = []
        base_metadata = metadata or {}
        base_metadata.update({"source": file_path, "file_type": "docx"})
        
        try:
            doc = DocxDocument(file_path)
            full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            
            # 提取表格内容
            for table in doc.tables:
                table_text = "\n[表格]\n"
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    table_text += row_text + "\n"
                full_text += table_text + "[/表格]\n"
            
            chunks = self.text_splitter.split_text(full_text)
            for i, chunk in enumerate(chunks):
                meta = base_metadata.copy()
                meta["chunk_index"] = i
                documents.append(Document(page_content=chunk, metadata=meta))
            
            logger.info(f"Word 加载完成: {file_path}, 共 {len(documents)} 个片段")
        except Exception as e:
            logger.error(f"Word 加载失败 {file_path}: {e}")
        
        return documents
    
    def _load_excel(self, file_path: str, metadata: Optional[dict] = None) -> List[Document]:
        """加载 Excel 文件"""
        documents = []
        base_metadata = metadata or {}
        base_metadata.update({"source": file_path, "file_type": "excel"})
        
        try:
            excel_file = pd.ExcelFile(file_path)
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                
                # 将DataFrame转换为文本描述
                sheet_content = f"工作表: {sheet_name}\n"
                sheet_content += f"列名: {', '.join(df.columns.astype(str))}\n"
                sheet_content += "数据内容:\n"
                
                # 只取前100行避免过大
                for idx, row in df.head(100).iterrows():
                    row_text = " | ".join([f"{col}: {val}" for col, val in row.items() if pd.notna(val)])
                    sheet_content += f"第{idx+1}行: {row_text}\n"
                
                meta = base_metadata.copy()
                meta["sheet_name"] = sheet_name
                documents.append(Document(page_content=sheet_content, metadata=meta))
            
            logger.info(f"Excel 加载完成: {file_path}, 共 {len(documents)} 个工作表")
        except Exception as e:
            logger.error(f"Excel 加载失败 {file_path}: {e}")
        
        return documents
    
    def _load_txt(self, file_path: str, metadata: Optional[dict] = None) -> List[Document]:
        """加载文本文件（支持多种编码）"""
        documents = []
        base_metadata = metadata or {}
        base_metadata.update({"source": file_path, "file_type": "txt"})
        
        # 尝试多种编码
        encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
        content = None
        
        for enc in encodings:
            try:
                with open(file_path, 'r', encoding=enc) as f:
                    content = f.read()
                break
            except UnicodeDecodeError:
                continue
        
        if content is None:
            logger.error(f"无法解码文本文件: {file_path}")
            return []
        
        chunks = self.text_splitter.split_text(content)
        for i, chunk in enumerate(chunks):
            meta = base_metadata.copy()
            meta["chunk_index"] = i
            documents.append(Document(page_content=chunk, metadata=meta))
        
        logger.info(f"TXT 加载完成: {file_path}, 共 {len(documents)} 个片段")
        return documents